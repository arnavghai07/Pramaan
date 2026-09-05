"""
reports/docx.py — PRAMAAN
==========================
The same ReportData as a real, editable Word document.

WHY A SEPARATE FORMAT AT ALL
-----------------------------
SIH26034 asks for an editable export, and a PDF is not one. An enforcement
officer's next step after an inspection is usually to add context the
system cannot know — what the shopkeeper said, which consignment the pack
came from, a reference to a previous notice — and then to put the document
into a departmental workflow that expects Word. So this produces a genuine
.docx: headings that appear in Word's navigation pane, real tables whose
rows can be edited or deleted, and normal paragraphs. Nothing is drawn as
a picture of text.

The content decisions are all in reports/builder.py. This file only places
the same strings the PDF places, which is what guarantees the two formats
of one inspection cannot disagree.

NOTE ON THE MODULE NAME: this is `reports.docx`, imported as such; the
`from docx import ...` below resolves to the top-level python-docx package
because Python 3 imports are absolute.
"""
from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from reports.builder import NOT_CAPTURED, ReportData
from reports.images import load_evidence

_STATUS_RGB = {
    "COMPLIANT": RGBColor(0x15, 0x80, 0x3D),
    "NON_COMPLIANT": RGBColor(0xB9, 0x1C, 0x1C),
    "NEEDS_MANUAL_REVIEW": RGBColor(0xB4, 0x53, 0x09),
}
_NEUTRAL = RGBColor(0x33, 0x41, 0x55)
_MUTED = RGBColor(0x47, 0x55, 0x69)

_STATE_RGB = {
    "PRESENT": RGBColor(0x15, 0x80, 0x3D),
    "MISSING": RGBColor(0xB9, 0x1C, 0x1C),
    "REVIEW": RGBColor(0xB4, 0x53, 0x09),
}

#: Usable width of an A4 page at the 1-inch margins python-docx defaults to.
#: Images are fitted to this so nothing is cropped by the page edge.
_CONTENT_WIDTH_IN = 6.3
_MAX_IMAGE_HEIGHT_IN = 7.5

#: Table style shipped with python-docx's default template. Guarded at use
#: because a document built from a customised template may not have it, and
#: a missing style raises rather than degrading.
_TABLE_STYLE = "Table Grid"


def _apply_table_style(table) -> None:
    try:
        table.style = _TABLE_STYLE
    except KeyError:
        pass    # borderless, but every value is still present and editable
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def _muted(document, text: str):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = _MUTED
    return paragraph


def _kv_table(document, pairs: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    _apply_table_style(table)
    for label, value in pairs:
        cells = table.add_row().cells
        cells[0].width = Inches(2.2)
        cells[1].width = Inches(_CONTENT_WIDTH_IN - 2.2)
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].text = value


def _rule6_table(document, data: ReportData) -> None:
    table = document.add_table(rows=1, cols=4)
    _apply_table_style(table)
    for cell, heading in zip(table.rows[0].cells,
                             ("Declaration", "State", "Extracted value",
                              "Requirement")):
        cell.paragraphs[0].add_run(heading).bold = True

    for row in data.rule6_rows:
        cells = table.add_row().cells
        cells[0].text = row.field
        state_run = cells[1].paragraphs[0].add_run(row.state)
        state_run.bold = True
        state_run.font.color.rgb = _STATE_RGB.get(row.state, _NEUTRAL)
        cells[2].text = row.value
        cells[3].text = row.requirement

    # Repeat the header row when the table spans pages, so a printed second
    # page is not four unlabelled columns.
    header = table.rows[0]
    try:
        header._tr.get_or_add_trPr().append(header._tr.makeelement(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblHeader",
            {}))
    except Exception:      # pragma: no cover - cosmetic only
        pass


def _bullets(document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def render_docx(data: ReportData) -> bytes:
    """The finished .docx for one stored inspection, as bytes."""
    document = Document()

    title = document.add_heading("PRAMAAN", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Legal Metrology (Packaged Commodities) "
                                    "compliance inspection report")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = _MUTED

    verdict = document.add_paragraph()
    verdict_run = verdict.add_run(f"OVERALL: {data.status_label}")
    verdict_run.bold = True
    verdict_run.font.size = Pt(16)
    verdict_run.font.color.rgb = _STATUS_RGB.get(data.overall_status, _NEUTRAL)
    if data.status_note:
        _muted(document, data.status_note)

    document.add_heading("Inspection record", level=1)
    _kv_table(document, [
        ("Inspection ID", f"#{data.inspection_id}"),
        ("Inspection date and time", data.inspected_at),
        ("Inspecting officer", data.inspector_name),
        ("Product", data.product_name),
        ("Manufacturer", data.manufacturer),
        ("Maximum retail price", data.mrp),
        ("Net quantity", data.net_quantity),
        ("Manufacture / packing date", data.mfg_date),
    ])

    document.add_heading("Rule 6 — mandatory declarations", level=1)
    document.add_paragraph(data.mandatory_summary)
    _rule6_table(document, data)

    document.add_paragraph()
    if data.missing_mandatory:
        document.add_heading("Missing or unresolved mandatory declarations",
                             level=2)
        _bullets(document, data.missing_mandatory)
    else:
        document.add_paragraph("No mandatory declaration was found missing.")

    if data.rule6_problems:
        document.add_heading("Cross-check problems", level=2)
        _bullets(document, data.rule6_problems)

    document.add_heading("Rule 7 — physical character height", level=1)
    _kv_table(document, data.rule7_rows)
    if data.rule7_note:
        _muted(document, data.rule7_note)

    document.add_heading("Additional compliance analysis", level=1)
    if data.analysis_note:
        _muted(document, data.analysis_note)
    for section in data.analysis_sections:
        suffix = " (advisory — does not affect the verdict)" if section.advisory else ""
        document.add_heading(
            f"{section.title} — {section.state_label}{suffix}", level=2)
        if section.explanation:
            document.add_paragraph(section.explanation)
        if section.findings:
            _bullets(document, section.findings)

    document.add_heading("Findings", level=1)
    if data.findings:
        _bullets(document, data.findings)
    else:
        # Same rule as the PDF: an empty section reads as "not checked".
        document.add_paragraph("No findings were recorded for this inspection.")

    # Evidence begins on a new page — the photographs are full width, and a
    # heading stranded above a page break separates a caption from its image.
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Evidence", level=1)
    for item in data.evidence:
        document.add_heading(item.caption, level=2)
        image = load_evidence(item.path)
        if image is None:
            _muted(document, NOT_CAPTURED)
            continue
        width_in, _ = image.size_for(_CONTENT_WIDTH_IN, _MAX_IMAGE_HEIGHT_IN)
        # Only the width is given: python-docx derives the height from the
        # image's own dimensions, which is what keeps the aspect ratio exact.
        paragraph = document.add_paragraph()
        paragraph.add_run().add_picture(image.stream(), width=Inches(width_in))

    document.add_paragraph()
    _muted(document,
           f"This document was generated by PRAMAAN on {data.generated_at}. "
           f"It reproduces inspection #{data.inspection_id} exactly as "
           f"recorded on {data.inspected_at}; no check was re-run and no "
           "verdict was recalculated to produce it.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
